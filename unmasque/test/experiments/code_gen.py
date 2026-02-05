import csv

from docx import Document
from docx.shared import Pt

# ---------------- CONFIG ----------------
CSV_FILE = "/Users/ahanapradhan/Desktop/exam_dj_sectionB.csv"
DOCX_FILE = "/Users/ahanapradhan/Desktop/exam_dj_sectionB.docx"
FONT_NAME = "Helvetica"
FONT_SIZE = 24
# ----------------------------------------


doc = Document()

with open(CSV_FILE, newline="", encoding="utf-8") as f:
    reader = csv.reader(f)
    header = next(reader, None)  # remove this line if no header

    first_record = True

    for row_num, row in enumerate(reader, start=1):
        if len(row) != 3:
            raise ValueError(f"Row {row_num} does not have exactly 4 columns")

        # Page break BEFORE every record except the first
        if not first_record:
            doc.add_page_break()
        first_record = False

        # -------- Heading --------
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run("EGC211P (Programming 2A) C++ GIE\nSection B\nDomjudge url: http://172.16.201.120\n")
        heading_run.font.name = "Helvetica"
        heading_run.font.size = Pt(24)

        table = doc.add_table(rows=3, cols=1)
        table.style = "Table Grid"

        for i, value in enumerate(row):
            cell = table.cell(i, 0)
            para = cell.paragraphs[0]
            para.clear()

            run = para.add_run(value)
            run.font.name = "Helvetica"
            run.font.size = Pt(18)

doc.save(DOCX_FILE)